import streamlit as st
from docx import Document
import io
import zipfile
import base64

def process_document(file, heading_level):
    doc = Document(file)
    sections = []
    current_section = []
    
    for para in doc.paragraphs:
        if para.style.name.lower().startswith('heading') and \
           (heading_level == 0 or para.style.name.lower() == f'heading {heading_level}' or 
            para.style.name.lower().startswith(f'heading {heading_level}')):
            if current_section:
                sections.append(current_section)
                current_section = []
        current_section.append(para)
    
    if current_section:
        sections.append(current_section)
    
    return sections

def create_docx(section):
    new_doc = Document()
    for para in section:
        new_para = new_doc.add_paragraph(para.text)
        new_para.style = new_doc.styles['Normal']
        if para.style.name.lower().startswith('heading'):
            new_para.style = new_doc.styles['Heading 1']
    
    doc_buffer = io.BytesIO()
    new_doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def get_docx_download_link(doc_buffer, filename):
    b64 = base64.b64encode(doc_buffer.getvalue()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-link"><i class="fas fa-download"></i></a>'

def render_document_content(section):
    content = ""
    for para in section:
        if para.style.name.lower().startswith('heading'):
            content += f"## {para.text}\n\n"
        else:
            content += f"{para.text}\n\n"
    return content

st.set_page_config(page_title="DOCX Processor", page_icon="📄", layout="wide")

st.markdown("""
<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.1/css/all.min.css">
<style>
    .file-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(100px, 1fr));
        gap: 10px;
        padding: 10px;
    }
    .file-item {
        display: flex;
        flex-direction: column;
        align-items: center;
        text-align: center;
        padding: 5px;
        border: 1px solid #ddd;
        border-radius: 5px;
        transition: all 0.3s;
    }
    .file-item:hover {
        box-shadow: 0 0 5px rgba(0,0,0,0.1);
    }
    .file-icon {
        font-size: 24px;
        color: #4CAF50;
        margin-bottom: 5px;
    }
    .file-name {
        font-size: 10px;
        word-wrap: break-word;
        max-width: 100%;
        margin-bottom: 5px;
    }
    .download-link {
        color: #4CAF50;
        text-decoration: none;
        margin-top: 3px;
    }
    .button-container {
        display: flex;
        justify-content: space-around;
        width: 100%;
        margin-top: 5px;
    }
    .content-box {
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 10px;
        max-height: 500px;
        overflow-y: auto;
    }
    .stButton > button {
        font-size: 10px;
        padding: 2px 5px;
        height: auto;
        min-height: 0px;
    }
</style>
""", unsafe_allow_html=True)

st.title("DOCX Processor")

# Initialize session state
if 'sections' not in st.session_state:
    st.session_state.sections = []
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None
if 'heading_choice' not in st.session_state:
    st.session_state.heading_choice = "Heading 1"
if 'selected_section' not in st.session_state:
    st.session_state.selected_section = None

uploaded_file = st.file_uploader("Choose a DOCX file", type="docx")

if uploaded_file is not None and uploaded_file != st.session_state.uploaded_file:
    st.session_state.uploaded_file = uploaded_file
    st.session_state.sections = []  # Reset sections when a new file is uploaded
    st.session_state.selected_section = None  # Reset selected section
    st.success("File successfully uploaded!")

heading_options = ["Heading 1", "Heading 2", "Heading 3", "Any Heading", "Page End"]
heading_choice = st.selectbox("Select division method:", heading_options, key='heading_choice')

if st.button("Process Document") and st.session_state.uploaded_file:
    heading_level = heading_options.index(heading_choice) + 1 if heading_choice != "Page End" else 0
    if heading_choice == "Any Heading":
        heading_level = 0
    
    with st.spinner("Processing document..."):
        st.session_state.sections = process_document(st.session_state.uploaded_file, heading_level)
    
    st.success(f"Document processed! Found {len(st.session_state.sections)} sections.")

# Function to handle deletion
def delete_section(index):
    st.session_state.sections.pop(index)
    st.session_state.selected_section = None  # Reset selected section after deletion

# Display file grid and content
if st.session_state.sections:
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown('<div class="file-grid">', unsafe_allow_html=True)
        for i, section in enumerate(st.session_state.sections):
            doc_buffer = create_docx(section)
            filename = f"Section_{i+1}.docx"
            preview = section[0].text[:20] + "..." if len(section[0].text) > 20 else section[0].text
            
            st.markdown(f"""
                <div class="file-item">
                    <i class="fas fa-file-word file-icon"></i>
                    <div class="file-name" title="{preview}">{filename}</div>
                    {get_docx_download_link(doc_buffer, filename)}
                </div>
            """, unsafe_allow_html=True)
            
            col1_1, col1_2 = st.columns(2)
            with col1_1:
                if st.button("View", key=f"view_{i}", help=f"View section {i+1}"):
                    st.session_state.selected_section = i
            with col1_2:
                if st.button("Delete", key=f"delete_{i}", help=f"Delete section {i+1}"):
                    delete_section(i)
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        if st.session_state.selected_section is not None:
            st.markdown("## Section Content")
            content = render_document_content(st.session_state.sections[st.session_state.selected_section])
            st.markdown(f"""
                <div class="content-box">
                    {content}
                </div>
            """, unsafe_allow_html=True)
    
    # Create ZIP file with remaining sections
    if st.session_state.sections:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
            for i, section in enumerate(st.session_state.sections, 1):
                doc_buffer = create_docx(section)
                zip_file.writestr(f"Section_{i}.docx", doc_buffer.getvalue())
        
        zip_buffer.seek(0)
        
        st.download_button(
            label="Download All Sections (ZIP)",
            data=zip_buffer,
            file_name="processed_sections.zip",
            mime="application/zip"
        )
else:
    st.info("Upload a .docx file, select a division method, and click 'Process Document' to begin.")
